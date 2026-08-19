# SPDX-FileCopyrightText: Copyright (c) 2026 NVIDIA CORPORATION & AFFILIATES. All rights reserved.
# SPDX-License-Identifier: Apache-2.0

"""Reviewable HTML, CSV, Office, and hash-receipted audit exports."""

from __future__ import annotations

import csv
import html
import json
import os
import shutil
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Iterable

from .chronology import ChronologyBuilder
from .schema import (
    EvidenceError,
    canonical_json,
    csv_safe,
    evidence_alias,
    safe_filename,
    sha256_file,
)
from .store import Workspace


class EvidenceExporter:
    def __init__(self, workspace: Workspace):
        self.workspace = workspace

    def export_all(
        self,
        case_id: str,
        output_directory: str | Path,
        *,
        include_originals: bool = True,
        draft: bool = False,
    ) -> dict[str, Any]:
        snapshot_id = self.workspace.head_snapshot_id(case_id)
        initial_verification = self.workspace.verify(case_id, snapshot_id)
        if not initial_verification["ok"]:
            raise EvidenceError(
                "cannot export an unverified case: " + "; ".join(initial_verification["failures"])
            )
        manifest = self.workspace.load_manifest(case_id, snapshot_id)
        review_anchor = ""
        for _attempt in range(3):
            review_anchor = self.workspace.verify_audit(case_id)["head"]
            chronology = ChronologyBuilder(self.workspace).build(
                case_id, snapshot_id, record_audit=False
            )
            verification = self.workspace.verify(case_id, snapshot_id)
            effective_coverage = self.workspace.effective_coverage(case_id, manifest)
            if self.workspace.verify_audit(case_id)["head"] == review_anchor:
                break
        else:
            raise EvidenceError("review state kept changing while export was prepared")
        pending_events = [
            event for event in chronology["events"] if event["review_status"] == "unreviewed"
        ]
        final_ready = effective_coverage["complete_for_negative_assertions"] and not pending_events
        if not draft and not final_ready:
            raise EvidenceError(
                "final export requires accepted review for every review-required page and "
                "a decision for every chronology event; use --draft for an explicitly "
                "marked review draft"
            )
        export_status = "final" if final_ready and not draft else "draft"

        requested = Path(output_directory).expanduser().absolute()
        requested.parent.mkdir(parents=True, mode=0o700, exist_ok=True)
        if requested.exists() or requested.is_symlink():
            raise EvidenceError(f"export destination must not already exist: {requested}")
        output = Path(tempfile.mkdtemp(prefix=".evidence-export.", dir=requested.parent))
        output.chmod(0o700)

        try:
            files: list[Path] = []
            snapshot_root = (
                self.workspace.root / "cases" / case_id / "snapshots" / manifest["snapshot_id"]
            )
            files.append(_copy_secure(snapshot_root / "manifest.json", output / "manifest.json"))
            files.append(
                _copy_secure(snapshot_root / "manifest.sha256", output / "manifest.sha256")
            )
            files.extend(self._copy_page_records(case_id, manifest, output))
            files.append(_write_json(output / "verification.json", verification))
            files.append(
                _write_json(
                    output / "coverage.json",
                    {
                        "snapshot_id": manifest["snapshot_id"],
                        "integrity_verified": verification["ok"],
                        "coverage": effective_coverage,
                    },
                )
            )
            files.append(_write_coverage_csv(output / "coverage.csv", manifest, verification["ok"]))
            files.append(
                _write_exclusions_csv(
                    output / "excluded-sources.csv",
                    effective_coverage.get("excluded_documents", []),
                    manifest["snapshot_id"],
                    verification["ok"],
                )
            )
            files.append(_write_json(output / "chronology.json", chronology))
            files.append(
                _write_chronology_csv(
                    output / "chronology.csv",
                    chronology["events"],
                    manifest["snapshot_id"],
                    verification["ok"],
                )
            )

            if include_originals:
                files.extend(self._copy_originals(case_id, manifest, output))
                files.extend(self._copy_page_images(case_id, manifest, output))

            files.append(
                _write_report_html(
                    output / "report.html",
                    case=self.workspace.get_case(case_id),
                    manifest=manifest,
                    verification=verification,
                    chronology=chronology,
                    coverage=effective_coverage,
                    export_status=export_status,
                )
            )
            files.append(
                _write_docx(
                    output / "report.docx",
                    manifest,
                    verification,
                    chronology,
                    effective_coverage,
                    export_status,
                )
            )
            files.append(
                _write_xlsx(
                    output / "chronology.xlsx",
                    chronology["events"],
                    manifest["snapshot_id"],
                    verification["ok"],
                    export_status,
                    effective_coverage,
                    manifest,
                )
            )

            audit_path = output / "audit.jsonl"
            audit_lines = [canonical_json(event) for event in self.workspace.audit_events(case_id)]
            _write_text_secure(audit_path, "\n".join(audit_lines) + ("\n" if audit_lines else ""))
            files.append(audit_path)

            artifact_hashes = {
                path.relative_to(output).as_posix(): {
                    "sha256": sha256_file(path),
                    "size_bytes": path.stat().st_size,
                }
                for path in sorted(files)
                if path.is_file()
            }
            bundle_receipt = {
                "schema": "evidence-workbench-export-v1",
                "integrity_kind": "unsigned-sha256-receipt",
                "case_id": case_id,
                "snapshot_id": manifest["snapshot_id"],
                "source_manifest_sha256": sha256_file(output / "manifest.json"),
                "audit_head": review_anchor,
                "export_status": export_status,
                "include_originals": include_originals,
                "files": artifact_hashes,
            }
            receipt_path = _write_json(output / "bundle-receipt.json", bundle_receipt)
            files.append(receipt_path)

            zip_path = output / "evidence-audit-bundle.zip"
            _write_deterministic_zip(zip_path, output, files)
            bundle_sha256 = sha256_file(zip_path)
            with self.workspace.publication_lock(case_id):
                if self.workspace.head_snapshot_id(case_id) != snapshot_id:
                    raise EvidenceError(
                        "evidence snapshot changed before export publication; retry the export"
                    )
                if self.workspace.verify_audit(case_id)["head"] != review_anchor:
                    raise EvidenceError(
                        "review state changed before export publication; retry the export"
                    )
                if requested.exists() or requested.is_symlink():
                    raise EvidenceError(f"export destination appeared concurrently: {requested}")
                os.rename(output, requested)
            result = {
                "case_id": case_id,
                "snapshot_id": manifest["snapshot_id"],
                "export_status": export_status,
                "output_directory": str(requested.resolve()),
                "audit_bundle": str((requested / zip_path.name).resolve()),
                "audit_bundle_sha256": bundle_sha256,
                "include_originals": include_originals,
                "files": sorted(path.relative_to(output).as_posix() for path in files),
            }
        except BaseException:
            if output.exists():
                shutil.rmtree(output, ignore_errors=True)
            raise
        self.workspace.record_event(
            case_id,
            "export_created",
            {
                "snapshot_id": manifest["snapshot_id"],
                "audit_bundle_sha256": result["audit_bundle_sha256"],
                "include_originals": include_originals,
                "export_status": export_status,
            },
        )
        return result

    def _copy_originals(self, case_id: str, manifest: dict[str, Any], output: Path) -> list[Path]:
        copied: list[Path] = []
        for source, excluded in _document_entries(manifest):
            source_hash = str(source["source_sha256"])
            filename = safe_filename(str(source["filename"]))
            relative = (
                Path("sources")
                / ("excluded" if excluded else "active")
                / str(source["document_id"])
                / filename
            )
            destination = output / relative
            destination.parent.mkdir(parents=True, mode=0o700, exist_ok=True)
            archived = self.workspace.source_path(case_id, source_hash)
            _copy_secure(archived, destination)
            if sha256_file(destination) != source_hash:
                destination.unlink(missing_ok=True)
                raise EvidenceError(f"exported source failed hash verification: {filename}")
            copied.append(destination)
        return copied

    def _copy_page_records(
        self, case_id: str, manifest: dict[str, Any], output: Path
    ) -> list[Path]:
        copied: list[Path] = []
        destination_root = output / "page_records"
        destination_root.mkdir(mode=0o700)
        seen: set[str] = set()
        for source, _excluded in _document_entries(manifest):
            for page in source["pages"]:
                record_hash = str(page["record_sha256"])
                if record_hash in seen:
                    continue
                seen.add(record_hash)
                source_path = (
                    self.workspace.root / "cases" / case_id / "page_records" / f"{record_hash}.json"
                )
                destination = destination_root / f"{record_hash}.json"
                _copy_secure(source_path, destination)
                if sha256_file(destination) != record_hash:
                    raise EvidenceError(f"exported page record failed verification: {record_hash}")
                copied.append(destination)
        return copied

    def _copy_page_images(self, case_id: str, manifest: dict[str, Any], output: Path) -> list[Path]:
        copied: list[Path] = []
        case_root = self.workspace.root / "cases" / case_id
        seen: set[str] = set()
        for source, _excluded in _document_entries(manifest):
            for page in source["pages"]:
                relative_name = str(page.get("evidence_image", ""))
                if not relative_name or relative_name in seen:
                    continue
                seen.add(relative_name)
                source_path = case_root / relative_name
                destination = output / relative_name
                destination.parent.mkdir(parents=True, mode=0o700, exist_ok=True)
                _copy_secure(source_path, destination)
                if sha256_file(destination) != page.get("evidence_image_sha256"):
                    destination.unlink(missing_ok=True)
                    raise EvidenceError(f"exported page image failed verification: {relative_name}")
                copied.append(destination)
        return copied


def _document_entries(
    manifest: dict[str, Any],
) -> Iterable[tuple[dict[str, Any], bool]]:
    for source in manifest["sources"]:
        yield source, False
    for tombstone in manifest.get("excluded_sources", []):
        yield tombstone["document"], True


def _write_json(path: Path, value: Any) -> Path:
    return _write_text_secure(
        path,
        json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
    )


def _write_coverage_csv(path: Path, manifest: dict[str, Any], integrity_verified: bool) -> Path:
    with path.open("w", encoding="utf-8", newline="") as stream:
        writer = csv.writer(stream)
        writer.writerow(
            [
                "snapshot_id",
                "integrity_verified",
                "source_alias",
                "document_id",
                "filename",
                "source_sha256",
                "page_number",
                "status",
                "extraction_method",
                "quality_score",
                "quality_score_kind",
                "needs_review",
                "error",
            ]
        )
        for source in manifest["sources"]:
            for page in source["pages"]:
                writer.writerow(
                    [
                        manifest["snapshot_id"],
                        integrity_verified,
                        csv_safe(
                            evidence_alias(str(source["filename"]), str(source["document_id"]))
                        ),
                        source["document_id"],
                        csv_safe(source["filename"]),
                        source["source_sha256"],
                        page["page_number"],
                        page["status"],
                        page["extraction_method"],
                        page["quality_score"],
                        page["quality_score_kind"],
                        page["needs_review"],
                        csv_safe(page["error"]),
                    ]
                )
    path.chmod(0o600)
    return path


def _write_chronology_csv(
    path: Path,
    events: Iterable[dict[str, Any]],
    snapshot_id: str,
    integrity_verified: bool,
) -> Path:
    columns = [
        "snapshot_id",
        "integrity_verified",
        "event_id",
        "raw_date",
        "normalized_date",
        "ambiguous",
        "normalization_note",
        "label",
        "document_id",
        "filename",
        "source_sha256",
        "page_number",
        "quote",
        "start_offset",
        "end_offset",
        "extraction_method",
        "needs_review",
        "review_status",
        "review_notes",
    ]
    with path.open("w", encoding="utf-8", newline="") as stream:
        writer = csv.DictWriter(stream, fieldnames=columns, extrasaction="ignore")
        writer.writeheader()
        for event in events:
            row = {column: csv_safe(event.get(column, "")) for column in columns}
            row["snapshot_id"] = snapshot_id
            row["integrity_verified"] = integrity_verified
            writer.writerow(row)
    path.chmod(0o600)
    return path


def _write_exclusions_csv(
    path: Path,
    excluded_documents: list[dict[str, Any]],
    snapshot_id: str,
    integrity_verified: bool,
) -> Path:
    columns = [
        "snapshot_id",
        "integrity_verified",
        "source_alias",
        "document_id",
        "filename",
        "source_sha256",
        "reason",
        "reviewer",
        "excluded_at",
        "tombstone_id",
    ]
    with path.open("w", encoding="utf-8", newline="") as stream:
        writer = csv.DictWriter(stream, fieldnames=columns)
        writer.writeheader()
        for excluded in excluded_documents:
            row = {column: csv_safe(excluded.get(column, "")) for column in columns}
            row["snapshot_id"] = snapshot_id
            row["integrity_verified"] = integrity_verified
            row["source_alias"] = csv_safe(
                evidence_alias(
                    str(excluded.get("filename", "")),
                    str(excluded.get("document_id", "")),
                )
            )
            writer.writerow(row)
    path.chmod(0o600)
    return path


def _write_report_html(
    path: Path,
    *,
    case: dict[str, Any],
    manifest: dict[str, Any],
    verification: dict[str, Any],
    chronology: dict[str, Any],
    coverage: dict[str, Any],
    export_status: str,
) -> Path:
    def esc(value: Any) -> str:
        return html.escape(str(value), quote=True)

    source_rows: list[str] = []
    for source in manifest["sources"]:
        filename = esc(source["filename"])
        alias = esc(evidence_alias(str(source["filename"]), str(source["document_id"])))
        source_path = f"sources/active/{esc(source['document_id'])}/{filename}"
        source_rows.append(
            "<tr>"
            f"<td>{alias}<br><small>Untrusted original: {source_path}</small></td>"
            f"<td><code>{esc(source['document_id'])}</code></td>"
            f"<td><code>{esc(source['source_sha256'])}</code></td>"
            f"<td>{esc(source['page_count'])}</td>"
            f"<td>{esc(source['status'])}</td>"
            "</tr>"
        )
    excluded_rows = [
        "<tr>"
        f"<td>{esc(evidence_alias(str(item['filename']), str(item['document_id'])))}</td>"
        f"<td><code>{esc(item['document_id'])}</code></td>"
        f"<td><code>{esc(item['source_sha256'])}</code></td>"
        f"<td>{esc(item['reason'])}</td>"
        f"<td>{esc(item['reviewer'])}</td>"
        f"<td>{esc(item['excluded_at'])}</td>"
        "</tr>"
        for item in coverage.get("excluded_documents", [])
    ]
    event_rows: list[str] = []
    for event in chronology["events"]:
        label = esc(event["label"])
        event_rows.append(
            "<tr>"
            f"<td>{esc(event['normalized_date'] or event['raw_date'])}</td>"
            f"<td>{label}</td>"
            f"<td>{esc(event['quote'])}</td>"
            f"<td>{esc(event['review_status'])}</td>"
            "</tr>"
        )
    warning = ""
    if not coverage["complete_for_negative_assertions"]:
        warning = (
            '<div class="warning"><strong>Coverage incomplete.</strong> '
            "Negative search results must not be interpreted as source-level absence."
            "</div>"
        )
    draft_warning = (
        '<div class="warning"><strong>DRAFT — human review incomplete.</strong> '
        "This export is not final work product.</div>"
        if export_status == "draft"
        else ""
    )
    document = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <meta http-equiv="Content-Security-Policy" content="default-src 'none'; style-src 'unsafe-inline'; img-src 'self' data:;">
  <title>Evidence report — {esc(case["name"])}</title>
  <style>
    body {{ font: 15px/1.5 system-ui, sans-serif; margin: 2rem auto; max-width: 1180px; color: #1b1b1b; padding: 0 1rem; }}
    h1, h2 {{ line-height: 1.2; }} code {{ word-break: break-all; font-size: .85em; }}
    table {{ border-collapse: collapse; width: 100%; margin: 1rem 0 2rem; }}
    th, td {{ border: 1px solid #ccc; padding: .55rem; text-align: left; vertical-align: top; }}
    th {{ background: #f3f3f3; }} .ok {{ color: #176b2c; }} .warning {{ border-left: 5px solid #b56b00; padding: .8rem; background: #fff5dd; }}
    .metrics {{ display: flex; flex-wrap: wrap; gap: .75rem; }} .metric {{ border: 1px solid #ccc; padding: .7rem 1rem; min-width: 8rem; }}
  </style>
</head>
<body>
  <h1>Evidence report: {esc(case["name"])}</h1>
  <p>Snapshot <code>{esc(manifest["snapshot_id"])}</code></p>
  <p>Export status: <strong>{esc(export_status.upper())}</strong></p>
  <p class="ok">Integrity verified: {esc(verification["ok"])}</p>
  {draft_warning}
  {warning}
  <h2>Coverage</h2>
  <div class="metrics">
    <div class="metric"><strong>{esc(coverage["documents_total"])}</strong><br>documents</div>
    <div class="metric"><strong>{esc(coverage["pages_total"])}</strong><br>pages</div>
    <div class="metric"><strong>{esc(coverage["pages_readable"])}</strong><br>readable</div>
    <div class="metric"><strong>{esc(len(coverage.get("review_pending_pages", [])))}</strong><br>pending review</div>
    <div class="metric"><strong>{esc(coverage["pages_failed"])}</strong><br>failed</div>
    <div class="metric"><strong>{esc(coverage.get("documents_excluded", 0))}</strong><br>excluded sources</div>
  </div>
  <h2>Sources</h2>
  <table><thead><tr><th>Source alias</th><th>Document ID</th><th>SHA-256</th><th>Pages</th><th>Status</th></tr></thead><tbody>{"".join(source_rows)}</tbody></table>
  <h2>Excluded sources</h2>
  <p>Excluded sources are not searched, but their authenticated tombstones and records remain in this export.</p>
  <table><thead><tr><th>Source alias</th><th>Document ID</th><th>SHA-256</th><th>Reason</th><th>Reviewer</th><th>Excluded at</th></tr></thead><tbody>{"".join(excluded_rows)}</tbody></table>
  <h2>Deterministic chronology</h2>
  <p>{esc(chronology["boundary"])}</p>
  <table><thead><tr><th>Date</th><th>Citation</th><th>Exact source context</th><th>Review</th></tr></thead><tbody>{"".join(event_rows)}</tbody></table>
</body>
</html>
"""
    return _write_text_secure(path, document)


def _write_docx(
    path: Path,
    manifest: dict[str, Any],
    verification: dict[str, Any],
    chronology: dict[str, Any],
    coverage: dict[str, Any],
    export_status: str,
) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise EvidenceError(
            "DOCX export requires python-docx; install the standalone application dependencies"
        ) from exc
    document = Document()
    document.add_heading("Evidence Workbench report", level=0)
    document.add_paragraph(f"Export status: {export_status.upper()}")
    document.add_paragraph(f"Snapshot: {manifest['snapshot_id']}")
    document.add_paragraph(f"Integrity verified: {verification['ok']}")
    document.add_heading("Coverage", level=1)
    document.add_paragraph(
        f"{coverage['documents_total']} documents; {coverage['pages_total']} pages; "
        f"{coverage['pages_readable']} readable; "
        f"{len(coverage.get('review_pending_pages', []))} pending review; "
        f"{coverage['pages_failed']} failed pages; "
        f"{coverage.get('documents_failed', 0)} failed documents; "
        f"{coverage.get('documents_excluded', 0)} excluded sources."
    )
    if not coverage["complete_for_negative_assertions"]:
        document.add_paragraph(
            "COVERAGE INCOMPLETE: negative search results do not prove source-level absence."
        )
    document.add_heading("Sources", level=1)
    for source in manifest["sources"]:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(
            evidence_alias(str(source["filename"]), str(source["document_id"]))
        ).bold = True
        paragraph.add_run(
            f" — document {source['document_id']}; source SHA-256 {source['source_sha256']}"
        )
    document.add_heading("Excluded sources", level=1)
    excluded_documents = coverage.get("excluded_documents", [])
    if not excluded_documents:
        document.add_paragraph("None.")
    for excluded in excluded_documents:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(
            evidence_alias(str(excluded["filename"]), str(excluded["document_id"]))
        ).bold = True
        paragraph.add_run(
            f" — document {excluded['document_id']}; {excluded['reason']} "
            f"(reviewer: {excluded['reviewer']}, "
            f"excluded: {excluded['excluded_at']})"
        )
    document.add_heading("Chronology", level=1)
    for event in chronology["events"]:
        date_value = event["normalized_date"] or event["raw_date"]
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(f"{date_value} {event['label']} — ").bold = True
        paragraph.add_run(event["quote"])
        paragraph.add_run(f" [{event['review_status'].upper()}]").italic = True
    document.save(path)
    path.chmod(0o600)
    return path


def _write_xlsx(
    path: Path,
    events: list[dict[str, Any]],
    snapshot_id: str,
    integrity_verified: bool,
    export_status: str,
    coverage: dict[str, Any],
    manifest: dict[str, Any],
) -> Path:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill
    except ImportError as exc:
        raise EvidenceError(
            "XLSX export requires openpyxl; install the standalone application dependencies"
        ) from exc
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Chronology"
    columns = [
        "normalized_date",
        "raw_date",
        "ambiguous",
        "label",
        "document_id",
        "quote",
        "filename",
        "source_sha256",
        "page_number",
        "needs_review",
        "review_status",
        "review_notes",
    ]
    sheet.append(columns)
    for cell in sheet[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="D9EAD3")
    for event in events:
        sheet.append([csv_safe(event.get(column, "")) for column in columns])
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    widths = {"A": 16, "B": 18, "C": 12, "D": 28, "E": 90, "F": 28, "G": 66}
    for column, width in widths.items():
        sheet.column_dimensions[column].width = width
    metadata = workbook.create_sheet("Metadata")
    metadata.append(["snapshot_id", snapshot_id])
    metadata.append(["integrity_verified", integrity_verified])
    metadata.append(["export_status", export_status])
    sources = workbook.create_sheet("Sources")
    source_columns = [
        "source_alias",
        "document_id",
        "filename",
        "source_sha256",
        "page_count",
        "status",
    ]
    sources.append(source_columns)
    for item in manifest["sources"]:
        source_row = {
            **item,
            "source_alias": evidence_alias(str(item["filename"]), str(item["document_id"])),
        }
        sources.append([csv_safe(source_row.get(column, "")) for column in source_columns])
    excluded = workbook.create_sheet("Excluded Sources")
    excluded_columns = [
        "source_alias",
        "document_id",
        "filename",
        "source_sha256",
        "reason",
        "reviewer",
        "excluded_at",
        "tombstone_id",
    ]
    excluded.append(excluded_columns)
    for item in coverage.get("excluded_documents", []):
        excluded_row = {
            **item,
            "source_alias": evidence_alias(
                str(item.get("filename", "")), str(item.get("document_id", ""))
            ),
        }
        excluded.append([csv_safe(excluded_row.get(column, "")) for column in excluded_columns])
    workbook.save(path)
    path.chmod(0o600)
    return path


def _write_deterministic_zip(zip_path: Path, root: Path, files: list[Path]) -> None:
    with zipfile.ZipFile(
        zip_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as archive:
        for path in sorted({candidate.resolve() for candidate in files}):
            if not path.is_file() or not path.is_relative_to(root.resolve()):
                continue
            relative = path.relative_to(root.resolve()).as_posix()
            info = zipfile.ZipInfo(relative, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            archive.writestr(info, path.read_bytes())
    zip_path.chmod(0o600)


def _copy_secure(source: Path, destination: Path) -> Path:
    if destination.exists() or destination.is_symlink():
        raise EvidenceError(f"refusing to overwrite export path: {destination}")
    destination.parent.mkdir(parents=True, mode=0o700, exist_ok=True)
    with source.open("rb") as input_stream, destination.open("xb") as output_stream:
        shutil.copyfileobj(input_stream, output_stream, length=1024 * 1024)
        output_stream.flush()
    destination.chmod(0o600)
    return destination


def _write_text_secure(path: Path, text: str) -> Path:
    if path.exists() or path.is_symlink():
        raise EvidenceError(f"refusing to overwrite export path: {path}")
    with path.open("x", encoding="utf-8", newline="\n") as stream:
        stream.write(text)
        stream.flush()
    path.chmod(0o600)
    return path
