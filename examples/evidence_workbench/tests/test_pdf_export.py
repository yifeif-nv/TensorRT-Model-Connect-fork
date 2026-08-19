# SPDX-FileCopyrightText: Copyright (c) 2026 NVIDIA CORPORATION & AFFILIATES. All rights reserved.
# SPDX-License-Identifier: Apache-2.0

"""PDF fallback and secure export tests.

Intent: Exercise the real PDFium/Pillow boundary and ensure exported HTML,
CSV, Office, and ZIP artifacts preserve hashes without executing source text.
Preconditions: Standalone application dependencies and a fake Model Connect OCR executable.
Postconditions: Scanned PDFs use OCR with review flags; exports escape markup,
neutralize formulas, contain original evidence, and carry a hash receipt.
"""

from __future__ import annotations

import csv
import json
import stat
import zipfile
from pathlib import Path

from PIL import Image, ImageDraw
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen.canvas import Canvas
import pytest

from evidence_workbench.chronology import ChronologyBuilder
import evidence_workbench.export as export_module
from evidence_workbench.export import EvidenceExporter
from evidence_workbench.ingest import Ingestor
from evidence_workbench.schema import EvidenceError, csv_safe, evidence_alias, sha256_file
from evidence_workbench.store import Workspace
from evidence_workbench.trtmc import TrtmcRunner


def _fake_trtmc(tmp_path: Path) -> tuple[Path, Path]:
    binary = tmp_path / "trtmc"
    bundle = tmp_path / "ocr.bundle"
    bundle.write_bytes(b"bundle")
    binary.write_text(
        """#!/usr/bin/env python3
import sys
if sys.argv[1] == 'inspect':
    print('family: deepseek_ocr')
elif sys.argv[1] == 'run':
    print('Scanned visit occurred on August 18, 2026. <script>alert(1)</script>')
else:
    raise SystemExit(2)
""",
        encoding="utf-8",
    )
    binary.chmod(0o755)
    return binary, bundle


def test_scanned_pdf_uses_model_connect_ocr(tmp_path: Path) -> None:
    image = Image.new("RGB", (640, 480), "white")
    ImageDraw.Draw(image).text((40, 80), "Scanned visit August 18 2026", fill="black")
    pdf = tmp_path / "scan.pdf"
    image.save(pdf, "PDF")
    image.close()
    binary, bundle = _fake_trtmc(tmp_path)
    workspace = Workspace(tmp_path / "workspace")
    workspace.create_case("PDF", "pdf")

    result = Ingestor(
        workspace,
        ocr_runner=TrtmcRunner(bundle, binary=binary),
        pdf_render_scale=1.2,
    ).ingest("pdf", pdf)

    page = result["snapshot"]["sources"][0]["pages"][0]
    assert page["extraction_method"] == "model_connect_ocr"
    assert page["status"] == "readable"
    assert page["needs_review"] is True
    assert page["evidence_image_sha256"]
    image_path = workspace.root / "cases/pdf" / page["evidence_image"]
    assert image_path.is_file()
    assert sha256_file(image_path) == page["evidence_image_sha256"]
    source = result["snapshot"]["sources"][0]
    workspace.record_review(
        "pdf",
        snapshot_id=result["snapshot"]["snapshot_id"],
        target_type="page",
        target_id=f"{source['document_id']}:p1",
        status="accepted",
        reviewer="Test Reviewer",
        expected_target_sha256=page["record_sha256"],
    )
    assert workspace.effective_coverage("pdf")["complete_for_negative_assertions"] is True


def test_mixed_native_text_and_image_pdf_uses_ocr(tmp_path: Path) -> None:
    body = Image.new("RGB", (500, 500), "white")
    ImageDraw.Draw(body).text((20, 30), "Scanned body contains hidden evidence", fill="black")
    body_path = tmp_path / "body.png"
    body.save(body_path, "PNG")
    body.close()
    pdf = tmp_path / "mixed.pdf"
    canvas = Canvas(str(pdf), pagesize=(612, 792))
    canvas.drawString(
        30, 760, "Native boilerplate header with more than forty printable characters"
    )
    canvas.drawImage(ImageReader(str(body_path)), 0, 0, width=612, height=700)
    canvas.save()
    binary, bundle = _fake_trtmc(tmp_path)
    workspace = Workspace(tmp_path / "workspace")
    workspace.create_case("Mixed", "mixed")

    result = Ingestor(workspace, ocr_runner=TrtmcRunner(bundle, binary=binary)).ingest("mixed", pdf)

    page = result["snapshot"]["sources"][0]["pages"][0]
    assert page["extraction_method"] == "model_connect_ocr"
    assert page["metadata"]["substantial_image_detected"] is True
    assert page["needs_review"] is True


def test_small_embedded_pdf_image_still_blocks_native_text_absence(tmp_path: Path) -> None:
    stamp = Image.new("RGB", (20, 20), "white")
    ImageDraw.Draw(stamp).text((1, 1), "X", fill="black")
    stamp_path = tmp_path / "stamp.png"
    stamp.save(stamp_path, "PNG")
    stamp.close()
    pdf = tmp_path / "small-image.pdf"
    canvas = Canvas(str(pdf), pagesize=(612, 792))
    canvas.drawString(
        30,
        760,
        "Native text is long and clean, but cannot cover a separate embedded image.",
    )
    canvas.drawImage(ImageReader(str(stamp_path)), 30, 700, width=10, height=10)
    canvas.save()
    binary, bundle = _fake_trtmc(tmp_path)
    workspace = Workspace(tmp_path / "workspace")
    workspace.create_case("Small image", "small-image")

    result = Ingestor(workspace, ocr_runner=TrtmcRunner(bundle, binary=binary)).ingest(
        "small-image", pdf
    )

    page = result["snapshot"]["sources"][0]["pages"][0]
    assert page["extraction_method"] == "model_connect_ocr"
    assert page["metadata"]["substantial_image_detected"] is True
    assert page["needs_review"] is True


def test_multiframe_tiff_preserves_frames_and_enforces_cap(tmp_path: Path) -> None:
    frames = [Image.new("RGB", (40, 30), color) for color in ("white", "gray")]
    tiff = tmp_path / "two-pages.tiff"
    frames[0].save(tiff, save_all=True, append_images=frames[1:], format="TIFF")
    for frame in frames:
        frame.close()
    binary, bundle = _fake_trtmc(tmp_path)

    workspace = Workspace(tmp_path / "workspace")
    workspace.create_case("TIFF", "tiff")
    result = Ingestor(workspace, ocr_runner=TrtmcRunner(bundle, binary=binary)).ingest("tiff", tiff)
    assert result["snapshot"]["coverage"]["pages_total"] == 2

    capped = Workspace(tmp_path / "capped")
    capped.create_case("Capped TIFF", "tiff")
    failed = Ingestor(
        capped,
        ocr_runner=TrtmcRunner(bundle, binary=binary),
        max_image_frames=1,
    ).ingest("tiff", tiff)
    assert failed["snapshot"]["coverage"]["documents_failed"] == 1
    assert failed["snapshot"]["coverage"]["complete_for_negative_assertions"] is False


def test_image_pixel_limit_fails_closed(tmp_path: Path) -> None:
    image = Image.new("RGB", (20, 20), "white")
    path = tmp_path / "large.png"
    image.save(path, "PNG")
    image.close()
    workspace = Workspace(tmp_path / "workspace")
    workspace.create_case("Pixels", "pixels")

    result = Ingestor(workspace, max_rendered_pixels=100).ingest("pixels", path)

    assert result["snapshot"]["coverage"]["documents_failed"] == 1
    assert result["snapshot"]["coverage"]["complete_for_negative_assertions"] is False


def test_export_is_escaped_reviewable_and_hash_receipted(tmp_path: Path) -> None:
    workspace = Workspace(tmp_path / "workspace")
    workspace.create_case("Case <script>alert(1)</script>", "export")
    source = tmp_path / "malicious.txt"
    source.write_text(
        "August 18, 2026. <script>alert('source')</script> =HYPERLINK('bad')",
        encoding="utf-8",
    )
    Ingestor(workspace).ingest("export", source)
    output = tmp_path / "out"

    result = EvidenceExporter(workspace).export_all("export", output, draft=True)

    report = (output / "report.html").read_text(encoding="utf-8")
    assert "<script>alert(1)</script>" not in report
    assert "&lt;script&gt;alert(1)&lt;/script&gt;" in report
    receipt = json.loads((output / "bundle-receipt.json").read_text(encoding="utf-8"))
    assert receipt["snapshot_id"] == result["snapshot_id"]
    assert receipt["integrity_kind"] == "unsigned-sha256-receipt"
    assert receipt["export_status"] == "draft"
    assert receipt["include_originals"] is True
    assert receipt["files"]["report.html"]["sha256"] == sha256_file(output / "report.html")
    with zipfile.ZipFile(result["audit_bundle"]) as archive:
        names = archive.namelist()
        assert "report.html" in names
        assert "report.docx" in names
        assert "chronology.xlsx" in names
        assert "manifest.sha256" in names
        assert any(name.startswith("page_records/") for name in names)
        assert any(name.startswith("sources/") for name in names)
        assert all(archive.getinfo(name).date_time == (1980, 1, 1, 0, 0, 0) for name in names)

    assert csv_safe("=2+2") == "'=2+2"
    with (output / "chronology.csv").open(encoding="utf-8", newline="") as stream:
        rows = list(csv.DictReader(stream))
    assert rows
    assert rows[0]["snapshot_id"] == result["snapshot_id"]
    assert rows[0]["quote"].startswith("August 18, 2026")
    assert workspace.verify("export")["ok"] is True
    original_manifest = (
        workspace.root / "cases/export/snapshots" / result["snapshot_id"] / "manifest.json"
    )
    assert (output / "manifest.json").read_bytes() == original_manifest.read_bytes()
    assert stat.S_IMODE(output.stat().st_mode) == 0o700
    for path in output.rglob("*"):
        if path.is_file():
            assert stat.S_IMODE(path.stat().st_mode) == 0o600


def test_final_export_requires_and_records_chronology_review(tmp_path: Path) -> None:
    workspace = Workspace(tmp_path / "workspace")
    workspace.create_case("Reviewed", "reviewed")
    source = tmp_path / "reviewed.txt"
    source.write_text("Visit occurred on August 18, 2026.", encoding="utf-8")
    result = Ingestor(workspace).ingest("reviewed", source)
    chronology = ChronologyBuilder(workspace).build("reviewed")
    event = chronology["events"][0]
    workspace.record_review(
        "reviewed",
        snapshot_id=result["snapshot"]["snapshot_id"],
        target_type="chronology_event",
        target_id=event["event_id"],
        status="accepted",
        reviewer="Test Reviewer",
        notes="Confirmed against source.",
    )

    exported = EvidenceExporter(workspace).export_all("reviewed", tmp_path / "final")

    assert exported["export_status"] == "final"


def test_export_refuses_existing_directory_and_child_symlink(tmp_path: Path) -> None:
    workspace = Workspace(tmp_path / "workspace")
    workspace.create_case("Safe export", "safe-export")
    source = tmp_path / "source.txt"
    source.write_text("No dated events here.", encoding="utf-8")
    Ingestor(workspace).ingest("safe-export", source)
    output = tmp_path / "existing"
    output.mkdir()
    victim = tmp_path / "victim.txt"
    victim.write_text("untouched", encoding="utf-8")
    (output / "manifest.json").symlink_to(victim)

    with pytest.raises(EvidenceError, match="must not already exist"):
        EvidenceExporter(workspace).export_all("safe-export", output)

    assert victim.read_text(encoding="utf-8") == "untouched"


def test_final_export_aborts_if_review_state_changes_before_publish(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    workspace = Workspace(tmp_path / "workspace")
    workspace.create_case("Race", "race")
    source = tmp_path / "race.txt"
    source.write_text("Visit occurred on August 18, 2026.", encoding="utf-8")
    ingested = Ingestor(workspace).ingest("race", source)
    event = ChronologyBuilder(workspace).build("race")["events"][0]
    workspace.record_review(
        "race",
        snapshot_id=ingested["snapshot"]["snapshot_id"],
        target_type="chronology_event",
        target_id=event["event_id"],
        status="accepted",
        reviewer="Reviewer",
    )
    original = export_module._write_report_html

    def revoke(*args: object, **kwargs: object) -> Path:
        workspace.record_review(
            "race",
            snapshot_id=ingested["snapshot"]["snapshot_id"],
            target_type="chronology_event",
            target_id=event["event_id"],
            status="unreviewed",
            reviewer="Reviewer",
        )
        return original(*args, **kwargs)  # type: ignore[arg-type]

    monkeypatch.setattr(export_module, "_write_report_html", revoke)
    output = tmp_path / "racy-export"

    with pytest.raises(EvidenceError, match="review state changed"):
        EvidenceExporter(workspace).export_all("race", output)

    assert not output.exists()


def test_excluded_source_remains_authenticated_and_visible_in_final_export(
    tmp_path: Path,
) -> None:
    from docx import Document
    from openpyxl import load_workbook

    workspace = Workspace(tmp_path / "workspace-exclusion")
    workspace.create_case("Exclusion", "exclusion")
    active = tmp_path / "active.txt"
    active.write_text("Active evidence without a dated event.", encoding="utf-8")
    failed = tmp_path / "wrong.bin"
    failed.write_bytes(b"mistaken unsupported upload")
    ingestor = Ingestor(workspace)
    ingestor.ingest("exclusion", active)
    failed_result = ingestor.ingest("exclusion", failed)
    failed_document = next(
        source
        for source in failed_result["snapshot"]["sources"]
        if source["filename"] == failed.name
    )

    manifest = workspace.exclude_document(
        "exclusion",
        failed_document["document_id"],
        reviewer="Test Reviewer",
        reason="Uploaded to the wrong matter.",
    )
    output = tmp_path / "excluded-export"
    exported = EvidenceExporter(workspace).export_all("exclusion", output)

    assert exported["export_status"] == "final"
    assert workspace.verify("exclusion")["ok"] is True
    assert manifest["coverage"]["documents_excluded"] == 1
    assert (
        manifest["excluded_sources"][0]["document"]["document_id"] == failed_document["document_id"]
    )
    report = (output / "report.html").read_text(encoding="utf-8")
    assert "Excluded sources" in report
    assert "Uploaded to the wrong matter." in report
    excluded_alias = evidence_alias(failed_document["filename"], failed_document["document_id"])
    assert excluded_alias in report
    with (output / "excluded-sources.csv").open(encoding="utf-8", newline="") as stream:
        rows = list(csv.DictReader(stream))
    assert rows[0]["document_id"] == failed_document["document_id"]
    assert rows[0]["source_alias"] == excluded_alias
    docx_text = "\n".join(
        paragraph.text for paragraph in Document(output / "report.docx").paragraphs
    )
    assert excluded_alias in docx_text
    workbook = load_workbook(output / "chronology.xlsx", read_only=True)
    excluded_rows = list(workbook["Excluded Sources"].iter_rows(values_only=True))
    assert excluded_rows[1][0] == excluded_alias
    excluded_original = (
        output / "sources" / "excluded" / failed_document["document_id"] / failed.name
    )
    assert excluded_original.read_bytes() == failed.read_bytes()


def test_export_disambiguates_sanitized_filename_collisions(tmp_path: Path) -> None:
    from docx import Document
    from openpyxl import load_workbook

    workspace = Workspace(tmp_path / "workspace-aliases")
    workspace.create_case("Aliases", "aliases")
    first = tmp_path / "same?.txt"
    second = tmp_path / "same*.txt"
    first.write_text("first active evidence", encoding="utf-8")
    second.write_text("second active evidence", encoding="utf-8")
    ingestor = Ingestor(workspace)
    ingestor.ingest("aliases", first)
    ingestor.ingest("aliases", second)
    manifest = workspace.load_manifest("aliases")
    assert {source["filename"] for source in manifest["sources"]} == {"same_.txt"}
    expected_aliases = {
        evidence_alias(str(source["filename"]), str(source["document_id"]))
        for source in manifest["sources"]
    }

    output = tmp_path / "alias-export"
    EvidenceExporter(workspace).export_all("aliases", output)

    report = (output / "report.html").read_text(encoding="utf-8")
    for source in manifest["sources"]:
        assert evidence_alias(source["filename"], source["document_id"]) in report
        assert source["document_id"] in report
    with (output / "coverage.csv").open(encoding="utf-8", newline="") as stream:
        coverage_rows = list(csv.DictReader(stream))
    assert {row["source_alias"] for row in coverage_rows} == expected_aliases
    assert {row["document_id"] for row in coverage_rows} == {
        source["document_id"] for source in manifest["sources"]
    }
    docx_text = "\n".join(
        paragraph.text for paragraph in Document(output / "report.docx").paragraphs
    )
    assert all(alias in docx_text for alias in expected_aliases)
    workbook = load_workbook(output / "chronology.xlsx", read_only=True)
    source_rows = list(workbook["Sources"].iter_rows(values_only=True))
    assert {str(row[0]) for row in source_rows[1:]} == expected_aliases
